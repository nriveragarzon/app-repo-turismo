# Liberias 
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.section import CT_SectPr
from docx.table import _Row

def set_cell_border(cell, **kwargs):
    """
    Establece los bordes de una celda.
    
    Uso:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
        left={"sz": 24, "val": "single", "color": "#0000FF", "space": "0"},
        right={"sz": 24, "val": "single", "color": "#000000", "space": "0"},
    )

    Args:
    cell: La celda a la que se aplicarán los bordes.
    kwargs: Un diccionario con las especificaciones para los bordes (tamaño, valor, color y espacio).
    """
    tcPr = cell._element.get_or_add_tcPr()

    # Verificar si ya existen bordes en la celda
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # Aplicar los bordes según las especificaciones proporcionadas
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))

def add_table_resumen(doc: Document, dataframe: pd.DataFrame, style: str, font_size: int, fuente: str):
    """
    Agrega una tabla al documento a partir de un DataFrame y asegura que no se divida entre páginas.
    También agrega una nota al final con la fuente de los datos, asegurando que esté en la misma página que la tabla.

    Args:
    doc (Document): El documento al que se añadirá la tabla.
    dataframe (DataFrame): El DataFrame que se convertirá en tabla.
    style (str): El estilo de la tabla.
    font_size (int): El tamaño de la letra para los títulos y el contenido de la tabla.
    fuente (str): La fuente de los datos.
    """
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        print(f"El valor proporcionado no es un DataFrame válido o está vacío: {dataframe}")
        return
    
    # Añadir la tabla al documento
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = doc.styles[style]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True  # Desactivar ajuste automático para controlar el ancho de las celdas
    
    # Configurar encabezados de la tabla
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column)
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_with_next = True  # Mantener la cabecera con la siguiente fila
        paragraph_format.keep_together = True  # Mantener la cabecera en una sola página
        hdr_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        set_cell_border(
            hdr_cells[i],
            top={"sz": 1, "val": "single", "color": "000000"},
            bottom={"sz": 1, "val": "single", "color": "000000"},
            left={"sz": 1, "val": "single", "color": "000000"},
            right={"sz": 1, "val": "single", "color": "000000"},
        )

    # Añadir filas de datos a la tabla
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            paragraph = row_cells[i].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_with_next = True  # Mantener cada celda con la siguiente en la misma fila
            paragraph_format.keep_together = True  # Mantener el contenido de la celda en una sola página
            row_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
            set_cell_border(
                row_cells[i],
                top={"sz": 1, "val": "single", "color": "000000"},
                bottom={"sz": 1, "val": "single", "color": "000000"},
                left={"sz": 1, "val": "single", "color": "000000"},
                right={"sz": 1, "val": "single", "color": "000000"},
            )

    # Aplicar formato especial a la primera fila (cabecera)
    for cell in hdr_cells:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "#215E99")
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Ajustar el ancho de las celdas para que la tabla ocupe todo el ancho de las márgenes del documento
    for row in table.rows:
        for cell in row.cells:
            cell.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # Agregar la fuente de los datos en un nuevo párrafo y asegurar que esté en la misma página que la tabla
    fuente_paragraph = doc.add_paragraph(f"Fuente: {fuente}", style='Normal')
    fuente_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    fuente_paragraph_format = fuente_paragraph.paragraph_format
    fuente_paragraph_format.keep_together = True  # Mantener la fuente en una sola página
    fuente_paragraph_format.keep_with_next = False  # No es necesario mantener con otro contenido después
    fuente_paragraph.space_before = Pt(12)  # Espacio antes del párrafo
    fuente_paragraph.space_after = Pt(12)  # Espacio después del párrafo
    fuente_paragraph.paragraph_format.left_indent = Cm(0.75)  # Indentación izquierda
    for run in fuente_paragraph.runs:
        run.font.size = Pt(9)  # Ajustar el tamaño de la fuente a 9 puntos
