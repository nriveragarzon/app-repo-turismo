# Librerías
from docx import Document
from io import BytesIO
import plotly.graph_objects as go
import pandas as pd
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK

# Importar módulos
from .style_objects import estilos
from .text_objects import add_heading, add_bullet_points
from .table_objects import add_table_resumen
from .utils import add_header_footer, agregar_tabla_contenidos, add_image_source

def transform_plotly_static(fig: go.Figure, image_format: str = "jpeg") -> bytes:
    """
    Convierte una figura de Plotly en una imagen estática en el formato especificado.

    Esta función verifica que el objeto proporcionado sea una instancia de
    `plotly.graph_objects.Figure`, actualiza su layout utilizando la plantilla
    'plotly_white' y lo convierte en una imagen estática utilizando Kaleido.

    Parámetros:
        fig (plotly.graph_objects.Figure): Objeto de figura de Plotly a convertir.
        image_format (str): Formato de la imagen resultante (por defecto "jpeg").
                            Se puede cambiar a otros formatos soportados por Kaleido (por ejemplo, "png").

    Retorna:
        bytes: Un objeto de tipo bytes que representa la imagen estática de la figura.
        None: Si el objeto proporcionado no es una figura válida de Plotly.
    """

    if isinstance(fig, go.Figure):
        # Actualiza el layout usando la plantilla 'plotly_white'
        fig.update_layout(template='plotly_white')
        # Convierte la figura a imagen en el formato especificado
        img_bytes = fig.to_image(format=image_format)
        return img_bytes
    else:
        return None
    

def documento_citi(dict_bullets, fig_time_series_viajeros, fig_time_series_gasto, fig_time_series_mice, fig_single_barchart_conectividad_mundo_frecuencias, fig_single_barchart_conectividad_colombia_frecuencias, fig_side_by_side_bar_gasto_promedio, df_resumen, pais_elegido, output_path, header_image_left, footer_image):

    """
    Genera un documento de Word con secciones de texto y gráficos relacionados con 
    el turismo internacional de un país específico hacia otros destinos (o desde 
    otros destinos hacia el país). El documento incluye una tabla de contenidos, 
    encabezados, pies de página y bullets descriptivos para cada indicador.

    Parámetros:
    -----------
    dict_bullets : dict
        Diccionario que contiene textos o viñetas (bullets) organizados por secciones 
        (p. ej., 'bullet_flujos_viajeros_mundo', 'bullet_medio_transporte', etc.). 
        Cada valor puede ser un texto tipo str o None en caso de no haber datos.
    fig_time_series_viajeros : plotly.graph_objs._figure.Figure o None
        Gráfico en formato Plotly que muestra la serie de tiempo del flujo de viajeros.
    fig_time_series_gasto : plotly.graph_objs._figure.Figure o None
        Gráfico en formato Plotly para la serie de tiempo del gasto promedio del viajero.
    fig_time_series_mice : plotly.graph_objs._figure.Figure o None
        Gráfico en formato Plotly para la serie de tiempo del flujo MICE (negocios).
    fig_single_barchart_conectividad_mundo_frecuencias : plotly.graph_objs._figure.Figure o None
        Gráfico de barras en formato Plotly, mostrando frecuencias aéreas internacionales 
        con el mundo.
    fig_single_barchart_conectividad_colombia_frecuencias : plotly.graph_objs._figure.Figure o None
        Gráfico de barras en formato Plotly, mostrando frecuencias aéreas internacionales 
        con Colombia.
    fig_side_by_side_bar_gasto_promedio : plotly.graph_objs._figure.Figure o None
        Gráfico en formato Plotly para comparar el gasto promedio con tarjeta de crédito 
        en distintos periodos o categorías.
    df_resumen : pandas.DataFrame
        DataFrame que contiene un resumen general (estadísticas, totales, etc.) de 
        las fuentes de datos consultadas.
    pais_elegido : str
        Nombre del país para el cual se generará el informe.
    output_path : str
        Ruta y nombre de archivo donde se guardará el documento final.
    header_image_left : str o bytes
        Imagen (ruta o bytes) para colocar en el encabezado izquierdo del documento.
    footer_image : str o bytes
        Imagen (ruta o bytes) para el pie de página del documento.

    Retorna:
    --------
    None
        La función no retorna ningún valor, sino que guarda el documento de Word 
        resultante en la ruta especificada por output_path.

    Detalles del proceso:
    ---------------------
    1. Se extraen las viñetas (bullets) relevantes desde el diccionario dict_bullets.
    2. Se convierten los objetos de gráficos de Plotly a imágenes en bytes si se proporcionan.
    3. Se agrupan las viñetas en listas por secciones (p. ej., indicadores de turismo al mundo, 
       gasto del viajero, conectividad aérea, etc.) y se filtran para remover valores None.
    4. Se crea el documento con python-docx, incluyendo:
       - Encabezado y pie de página personalizados (header_image_left, footer_image).
       - Título principal con el nombre del país.
       - Tabla de contenido, actualizable al abrir el archivo en Word (requiere presionar F9).
       - Secciones con encabezados, gráficos insertados y bullets descriptivos.
       - Tabla con el resumen de datos (df_resumen), si está disponible.
       - Un disclaimer al final, con formato de negrita, para aclarar la naturaleza 
         referencial de la información.
    5. Se guarda el documento en la ubicación especificada (output_path).

    Nota importante:
    ----------------
    - Para que la tabla de contenidos muestre la numeración de páginas correctamente, 
      se debe abrir el documento en Microsoft Word y actualizar manualmente los campos 
      (usualmente Ctrl + A y luego F9).
    - Algunos estilos adicionales de Word (fuentes, negritas, subrayados, etc.) 
      pueden requerir que el usuario tenga dichas fuentes instaladas o que edite 
      manualmente en Word para ajustar la apariencia.
    """
    
    # Extraer bullets
    bullet_flujos_viajeros_mundo = dict_bullets.get('bullet_flujos_viajeros_mundo', None)
    bullet_medio_transporte = dict_bullets.get('bullet_medio_transporte', None)
    bullet_noches_percnotacion = dict_bullets.get('bullet_noches_percnotacion', None)
    bullet_rango_edad = dict_bullets.get('bullet_rango_edad', None)
    bullet_motivo_viaje = dict_bullets.get('bullet_motivo_viaje', None)
    bullet_forma_viaje = dict_bullets.get('bullet_forma_viaje', None)
    bullet_destinos_internacionales_t_1 = dict_bullets.get('bullet_destinos_internacionales_t_1', None)
    bullet_destinos_internacionales_t = dict_bullets.get('bullet_destinos_internacionales_t', None)
    bullet_gasto_promedio = dict_bullets.get('bullet_gasto_promedio', None)
    bullet_gasto_categoria = dict_bullets.get('bullet_gasto_categoria', None)
    bullet_mice = dict_bullets.get('bullet_mice', None)
    bullet_frecuencias_mundo = dict_bullets.get('bullet_frecuencias_mundo', None)
    bullet_paises_con_frecuencias = dict_bullets.get('bullet_paises_con_frecuencias', None)
    bullet_frecuencias_destino_cerrado_t = dict_bullets.get('bullet_frecuencias_destino_cerrado_t', None)
    bullet_reservas_aereas_mex_cost_chi_per = dict_bullets.get('bullet_reservas_aereas_mex_cost_chi_per', None)
    bullet_busquedas_aereas_mex_cost_chi_per = dict_bullets.get('bullet_busquedas_aereas_mex_cost_chi_per', None)
    bullet_frecuencias_colombia = dict_bullets.get('bullet_frecuencias_colombia', None)
    bullet_frecuencias_municipio_cerrado_t = dict_bullets.get('bullet_frecuencias_municipio_cerrado_t', None)
    bullet_gasto_credibanco_cerrado_promedio = dict_bullets.get('bullet_gasto_credibanco_cerrado_promedio', None)
    bullet_gasto_directo_indirecto_credibanco_cerrado = dict_bullets.get('bullet_gasto_directo_indirecto_credibanco_cerrado', None)
    bullet_gasto_directo_cerrado = dict_bullets.get('bullet_gasto_directo_cerrado', None)
    bullet_gasto_indirecto_cerrado = dict_bullets.get('bullet_gasto_indirecto_cerrado', None)
    bullet_reservas_aereas_colombia = dict_bullets.get('bullet_reservas_aereas_colombia', None)
    bullet_busquedas_aereas_colombia = dict_bullets.get('bullet_busquedas_aereas_colombia', None)

    # Obtener bytes de los gráficos
    if isinstance(fig_time_series_viajeros, go.Figure):
        bytes_fig_time_series_viajeros = BytesIO(transform_plotly_static(fig = fig_time_series_viajeros, image_format = "jpeg"))
    else:
        bytes_fig_time_series_viajeros = None

    if isinstance(fig_time_series_gasto, go.Figure):
        bytes_fig_time_series_gasto = BytesIO(transform_plotly_static(fig = fig_time_series_gasto, image_format = "jpeg"))
    else:
        bytes_fig_time_series_gasto = None

    if isinstance(fig_time_series_mice, go.Figure):
        bytes_fig_time_series_mice = BytesIO(transform_plotly_static(fig = fig_time_series_mice, image_format = "jpeg"))
    else:
        bytes_fig_time_series_mice = None

    if isinstance(fig_single_barchart_conectividad_mundo_frecuencias, go.Figure):
        bytes_fig_single_barchart_conectividad_mundo_frecuencias = BytesIO(transform_plotly_static(fig = fig_single_barchart_conectividad_mundo_frecuencias, image_format = "jpeg"))
    else:
        bytes_fig_single_barchart_conectividad_mundo_frecuencias = None

    if isinstance(fig_single_barchart_conectividad_colombia_frecuencias, go.Figure):
        bytes_fig_single_barchart_conectividad_colombia_frecuencias = BytesIO(transform_plotly_static(fig = fig_single_barchart_conectividad_colombia_frecuencias, image_format = "jpeg"))
    else:
        bytes_fig_single_barchart_conectividad_colombia_frecuencias = None

    if isinstance(fig_side_by_side_bar_gasto_promedio, go.Figure):
        bytes_fig_side_by_side_bar_gasto_promedio = BytesIO(transform_plotly_static(fig = fig_side_by_side_bar_gasto_promedio, image_format = "jpeg"))
    else:
        bytes_fig_side_by_side_bar_gasto_promedio = None

    # Crear lista de bullets por sección

    ###################################
    # Flujos de viajeros hacia el mundo
    ###################################

    bullet_destinos_internacionales = None

    # Si al menos uno de los dos bullets tiene datos
    if bullet_destinos_internacionales_t or bullet_destinos_internacionales_t_1:
        # Unir solo los bullets que sean distintos de None
        parts = []
        if bullet_destinos_internacionales_t is not None:
            parts.append(bullet_destinos_internacionales_t)
        if bullet_destinos_internacionales_t_1 is not None:
            parts.append(bullet_destinos_internacionales_t_1)

        # Si al final hay algo que unir, crear la cadena
        if parts:
            bullet_destinos_internacionales = " ".join(parts)

    # Crear la lista (incluyendo el nuevo bullet_destinos_internacionales)
    bullets_seccion_flujo_viajeros_mundo = [
        bullet_flujos_viajeros_mundo,
        bullet_medio_transporte,
        bullet_noches_percnotacion,
        bullet_rango_edad,
        bullet_motivo_viaje,
        bullet_forma_viaje,
        bullet_destinos_internacionales
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_flujo_viajeros_mundo = [b for b in bullets_seccion_flujo_viajeros_mundo if b is not None]

    #################################
    # Gasto internacional del viajero
    #################################

    # Crear la lista
    bullets_seccion_gasto_internacionales_viajero = [
        bullet_gasto_promedio,
        bullet_gasto_categoria
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_gasto_internacionales_viajero = [b for b in bullets_seccion_gasto_internacionales_viajero if b is not None]

    ######
    # MICE
    ######

    # Crear la lista
    bullets_seccion_mice = [
        bullet_mice
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_mice = [b for b in bullets_seccion_mice if b is not None]

    ###########################
    # Conectividad con el mundo
    ###########################

    # Crear la lista
    bullets_seccion_conectividad_mundo = [
        bullet_frecuencias_mundo,
        bullet_paises_con_frecuencias,
        bullet_frecuencias_destino_cerrado_t
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_conectividad_mundo = [b for b in bullets_seccion_conectividad_mundo if b is not None]

    #############################################################
    # Reservas y búsquedas hacia México, Costa Rica, Chile y Perú
    #############################################################

    # Crear la lista
    bullets_seccion_res_bus_mex_cr_ch_per = [
        bullet_reservas_aereas_mex_cost_chi_per,
        bullet_busquedas_aereas_mex_cost_chi_per
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_res_bus_mex_cr_ch_per = [b for b in bullets_seccion_res_bus_mex_cr_ch_per if b is not None]

    ###########################
    # Conectividad con Colombia
    ###########################

    # Crear la lista
    bullets_seccion_conectividad_colombia = [
        bullet_frecuencias_colombia,
        bullet_frecuencias_municipio_cerrado_t
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_conectividad_colombia = [b for b in bullets_seccion_conectividad_colombia if b is not None]

    ##############################
    # Gasto con tarjeta de crédito
    ##############################

    # Crear la lista
    bullets_seccion_gasto_tarjeta_credito = [
        bullet_gasto_credibanco_cerrado_promedio,
        bullet_gasto_directo_indirecto_credibanco_cerrado,
        bullet_gasto_directo_cerrado,
        bullet_gasto_indirecto_cerrado
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_gasto_tarjeta_credito = [b for b in bullets_seccion_gasto_tarjeta_credito if b is not None]

    #####################################
    # Reservas y búsquedas hacia Colombia
    #####################################

    # Crear la lista
    bullets_seccion_res_bus_colombia = [
        bullet_reservas_aereas_colombia,
        bullet_busquedas_aereas_colombia
    ]

    # Filtrar valores None para que la lista contenga solo bullets válidos
    bullets_seccion_res_bus_colombia = [b for b in bullets_seccion_res_bus_colombia if b is not None]

    # Fuentes
    global_data_fuente = 'GlobalData'
    oag_fuente = 'OAG'
    fk_fuente = 'ForwardKeys'
    credibanco_fuente = 'Credibanco'
    iata_fuente = 'IATA-GAP'
    iata_nota = 'Datos de 2024 actualizados al tercer trimestre'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Crear documento
    doc = Document()

    # Agregar el footer header
    add_header_footer(doc, header_image_left=header_image_left, footer_image=footer_image)

    # Aplicar los estilos personalizados al documento
    estilos(doc)

    # Agregar un título principal
    add_heading(doc, text=f"Centro Internacional de Turismo Internacional: {pais_elegido}", level=0, style='Title')

    #########
    # Resumen
    #########
    if (not df_resumen.empty):

        # Título
        add_heading(doc, 'Resultados Generales', level=2, style='Heading 1')

        # Resumen
        add_table_resumen(doc, pd.DataFrame(df_resumen), 'Table Grid', 10, "GlobalData, OAG, Credibanco y IATA-GAP")

        # Salto de página 
        doc.add_page_break()


    ####################
    # Tabla de contenido
    ####################
    agregar_tabla_contenidos(doc)
    doc.add_page_break()


    #######################################
    # Indicadores de turismo hacia el mundo
    #######################################
    add_heading(doc, f'Indicadores de turismo de {pais_elegido} hacia el mundo', level=2, style='Heading 1')

    ###################################
    # Flujos de viajeros hacia el mundo
    ###################################
    if bullets_seccion_flujo_viajeros_mundo:

        # Título
        add_heading(doc, f'Flujo de viajeros de {pais_elegido} hacia el mundo', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_time_series_viajeros:
            add_image_source(doc=doc, img_bytes=bytes_fig_time_series_viajeros, fuente=global_data_fuente, titulo=f"Flujo de viajeros (miles) de {pais_elegido} hacia el mundo", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_flujo_viajeros_mundo)

    #################################
    # Gasto internacional del viajero
    #################################
    if bullets_seccion_gasto_internacionales_viajero:

        # Título
        add_heading(doc, f'Gasto internacional del viajero proveniente de {pais_elegido}', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_time_series_gasto:
            add_image_source(doc, bytes_fig_time_series_gasto, global_data_fuente, titulo=f"Gasto promedio del viajero de {pais_elegido} al mundo", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_gasto_internacionales_viajero)

    ######
    # MICE
    ######
    if bullets_seccion_mice:

        # Título
        add_heading(doc, f'Flujo de viajeros internacionales de {pais_elegido} por motivo de negocios', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_time_series_mice:
            add_image_source(doc, bytes_fig_time_series_mice, global_data_fuente, titulo=f"Flujo de viajeros internacionales de {pais_elegido} por motivo de Reuniones, incentivos, congresos y exposiciones (MICE)", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_mice)

    ###########################
    # Conectividad con el mundo
    ###########################
    if bullets_seccion_conectividad_mundo:

        # Título
        add_heading(doc, f'Conectividad aérea directa de {pais_elegido} con el mundo', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_single_barchart_conectividad_mundo_frecuencias:
            add_image_source(doc, bytes_fig_single_barchart_conectividad_mundo_frecuencias, oag_fuente, titulo=f"Conectividad de {pais_elegido} con el mundo: Frecuencias", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_conectividad_mundo)

    #############################################################
    # Reservas y búsquedas hacia México, Costa Rica, Chile y Perú
    #############################################################
    if bullets_seccion_res_bus_mex_cr_ch_per:

        # Título
        add_heading(doc, f'Reservas y búsquedas aéreas de {pais_elegido} hacia México, Costa Rica, Chile y Perú', level=3, style='Heading 2')

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_res_bus_mex_cr_ch_per)

    #######################################
    # Indicadores de turismo hacia Colombia
    #######################################
    add_heading(doc, f'Indicadores de turismo de {pais_elegido} hacia Colombia', level=2, style='Heading 1')

    ###########################
    # Conectividad con Colombia
    ###########################
    if bullets_seccion_conectividad_colombia:

        # Título
        add_heading(doc, f'Conectividad aérea directa de {pais_elegido} con Colombia', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_single_barchart_conectividad_colombia_frecuencias:
            add_image_source(doc, bytes_fig_single_barchart_conectividad_colombia_frecuencias, oag_fuente, titulo=f"Conectividad de {pais_elegido} con Colombia: Frecuencias", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_conectividad_colombia)

    ##############################
    # Gasto con tarjeta de crédito
    ##############################
    if bullets_seccion_gasto_tarjeta_credito:

        # Título
        add_heading(doc, f'Gasto con tarjeta de crédito en Colombia de los viajeros de {pais_elegido}', level=3, style='Heading 2')

        # Gráfico
        if bytes_fig_side_by_side_bar_gasto_promedio:
            add_image_source(doc, bytes_fig_side_by_side_bar_gasto_promedio, credibanco_fuente, titulo=f"Gasto promedio con tarjeta de crédito de los viajeros de {pais_elegido}", width_inches = 4.5)

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_gasto_tarjeta_credito)


    #####################################
    # Reservas y búsquedas hacia Colombia
    #####################################
    if bullets_seccion_res_bus_colombia:

        # Título
        add_heading(doc, f'Reservas y búsquedas aéreas de {pais_elegido} hacia Colombia', level=3, style='Heading 2')

        # Agregar bullets
        add_bullet_points(doc, bullets_seccion_res_bus_colombia)


    ############
    # Disclaimer
    ############

    # Salto de página
    doc.add_page_break()

    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)

    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(output_path)