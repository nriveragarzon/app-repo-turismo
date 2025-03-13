# Librerías
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO


def add_header_footer(doc: Document, header_image_left: str, footer_image: str, footer_text: str):
    """
    Agrega un encabezado y un pie de página a la primera sección de un documento de Word.
    
    En el encabezado se inserta una tabla con una sola fila y dos columnas:
    - Imagen en la celda izquierda con dimensiones personalizadas (7.66 cm de ancho, 3.4 cm de alto).
    - La celda derecha se mantiene vacía (podrías ampliarla según tus necesidades).

    En el pie de página se inserta una tabla con una sola fila y tres columnas:
    - Texto en la celda izquierda (alineado a la izquierda).
    - Número de página en la celda central (alineado al centro).
    - Imagen en la celda derecha (alineada a la derecha), con dimensiones personalizadas (5.99 cm de ancho, 1.27 cm de alto).

    Args:
        doc (Document): El objeto Document al que se añadirán encabezado y pie de página.
        header_image_left (str): Ruta de la imagen que se ubicará en la parte izquierda del encabezado.
        footer_image (str): Ruta de la imagen que se colocará en la parte derecha del pie de página.
        footer_text (str): Texto que aparecerá en la parte izquierda del pie de página.
    """
    # Obtener la primera sección del documento
    section = doc.sections[0]
    # Ajustar la distancia del encabezado y el pie de página
    section.header_distance = Cm(2.57)
    section.footer_distance = Cm(0.51)

 # Encabezado
    header = section.header
    #header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    #header_table.autofit = True

    # Imagen izquierda en el encabezado
    #header_cell_left = header_table.cell(0, 0)
    header_paragraph_left = header.paragraphs[0]
    header_run_left = header_paragraph_left.add_run()
    header_run_left.add_picture(header_image_left, width=Inches(2.5))

    # Pie de página
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=doc.sections[0].page_width)
    footer_table.autofit = True

    # Texto izquierda en el pie de página
    footer_cell_left = footer_table.cell(0, 0)
    footer_paragraph_left = footer_cell_left.paragraphs[0]
    footer_paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_run_left = footer_paragraph_left.add_run("Calle 30 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co")
    footer_run_left.font.size = Pt(8)

    # Número de página en el centro del pie de página
    footer_cell_center = footer_table.cell(0, 1)
    footer_paragraph_center = footer_cell_center.paragraphs[0]
    footer_paragraph_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run_center = footer_paragraph_center.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run_center._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    footer_run_center._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run_center._r.append(fldChar2)
    # Establecer el tamaño de la fuente del número de página
    footer_run_center.font.size = Pt(8)

    # Imagen derecha en el pie de página
    footer_cell_right = footer_table.cell(0, 2)
    footer_paragraph_right = footer_cell_right.paragraphs[0]
    footer_paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run_right = footer_paragraph_right.add_run()
    footer_run_right.add_picture(footer_image, width=Inches(2.0))

def agregar_tabla_contenidos(new_doc, font_size=8):
    """
    Agrega una tabla de contenidos al documento proporcionado en el estilo 'Normal'.
    
    Args:
    new_doc (Document): El documento al que se añadirá la tabla de contenidos.
    font_size (int): Tamaño de la fuente para todo el documento en puntos (por defecto 12).
    """
    # Configurar el estilo de fuente y tamaño para todo el documento
    style = new_doc.styles['Normal']
    font = style.font
    font.size = Pt(font_size)

    # Crear un párrafo para el título de la tabla de contenidos y centrarlo
    para = new_doc.add_paragraph("Tabla de Contenidos")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Century Gothic'
        run.font.size = Pt(11)
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0, 32, 96)

    # Crear un párrafo vacío para insertar el campo de la tabla de contenidos
    paragraph = new_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()
    
    # Insertar el campo de la tabla de contenidos
    fldChar = OxmlElement('w:fldChar')  # crea un nuevo elemento
    fldChar.set(qn('w:fldCharType'), 'begin')  # establece el atributo en el elemento
 
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # establece el atributo en el elemento
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # cambia 1-3 dependiendo de los niveles de encabezado que necesites
 
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
 
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Haga clic derecho para actualizar el campo."
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    fldChar2.append(fldChar3)
 
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
 
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
 
    p_element = paragraph._p

def add_image_source(doc: Document, img_bytes: BytesIO, fuente: str, titulo: str = None, width_inches: float = 6.5, height_inches: float = None) -> None:
    """
    Inserta una imagen, un título y su fuente asociada en un documento Word, 
    centrando todos los elementos y manteniéndolos juntos en la misma página.

    La función realiza lo siguiente:
      1. Inserta un párrafo con el título de la imagen (parámetro 'titulo') encima de la imagen. 
         Este párrafo se centra, tiene un tamaño de fuente de 8 puntos y se formatea para que permanezca
         junto al siguiente elemento (la imagen).
      2. Inserta la imagen con el tamaño especificado (manteniendo la relación de aspecto si no se indica la altura)
         y centra el párrafo donde se inserta la imagen, configurándolo para que se mantenga unido con el siguiente párrafo.
      3. Agrega un párrafo con la fuente (parámetro 'fuente') justo debajo de la imagen, centrado y con el mismo tamaño
         de fuente (8 puntos), formateado para mantenerse en la misma página que la imagen.

    Parámetros:
        doc (Document): Objeto de documento Word (python‑docx) donde se insertará la imagen.
        img_bytes (BytesIO): Objeto BytesIO de la imagen a insertar.
        fuente (str): Texto que indica la fuente de la imagen o de los datos.
        titulo (str): Título de la imagen que se mostrará encima de la misma. (Opcional)
        width_inches (float): Ancho de la imagen en pulgadas. Por defecto es 6.5.
        height_inches (float, opcional): Altura de la imagen en pulgadas. Si se omite, se mantiene la proporción original.

    Retorna:
        None: La función modifica el documento `doc` directamente.

    Ejemplo de uso:
        from docx import Document
        from io import BytesIO
        # Supongamos que 'img_bytes' es un objeto BytesIO obtenido a partir de una imagen.
        doc = Document()
        add_image_source(doc, img_bytes, "Fuente: Datos de XYZ", "Título de la Imagen", width_inches=6.5)
        doc.save("documento_con_imagen.docx")
    """
    # Verificar que img_bytes sea del tipo esperado
    if not isinstance(img_bytes, BytesIO):
        print("El objeto proporcionado no es de tipo 'BytesIO'.")
        return

    # Agregar el título de la imagen encima de la imagen
    if titulo:
        title_paragraph = doc.add_paragraph(titulo, style='Normal')
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_paragraph.paragraph_format.keep_together = True
        title_paragraph.paragraph_format.keep_with_next = True  # Mantener junto al siguiente elemento (imagen)
        for run in title_paragraph.runs:
            run.font.size = Pt(8)
            run.font.bold = True

    # Insertar la imagen en el documento con el tamaño especificado.
    if height_inches:
        doc.add_picture(img_bytes, width=Inches(width_inches), height=Inches(height_inches))
    else:
        doc.add_picture(img_bytes, width=Inches(width_inches))

    # Centrar la imagen obteniendo el último párrafo (donde se insertó la imagen) y configurando su alineación
    if doc.paragraphs:
        image_paragraph = doc.paragraphs[-1]
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Configurar para que la imagen se mantenga junto al siguiente elemento (la fuente)
        image_paragraph.paragraph_format.keep_with_next = True

    # Agregar un párrafo para la fuente debajo de la imagen
    fuente_paragraph = doc.add_paragraph(f"Fuente: {fuente}", style='Normal')
    fuente_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fuente_paragraph.paragraph_format.keep_together = True
    fuente_paragraph.paragraph_format.keep_with_next = False
    fuente_paragraph.space_before = Pt(12)
    fuente_paragraph.space_after = Pt(12)
    fuente_paragraph.paragraph_format.left_indent = Cm(0.75)
    for run in fuente_paragraph.runs:
        run.font.size = Pt(8)
        run.font.bold = True
