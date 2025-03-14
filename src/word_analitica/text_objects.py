# Librerías
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_heading(doc: Document, text: str, level: int, style: str = None):
    """
    Agrega un encabezado al documento con un estilo y nivel personalizados.

    Args:
    doc (Document): El documento al que se añadirá el encabezado.
    text (str): El texto del encabezado.
    level (int): El nivel del encabezado (0 para título, 1-9 para encabezados jerárquicos).
    style (str, optional): El estilo del encabezado (por defecto usa el estilo predeterminado del nivel).
    """
    heading = doc.add_heading(text, level=level)
    if style:
        heading.style = doc.styles[style]


def add_paragraph(doc: Document, text: str, style: str = None, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    Agrega un párrafo al documento con estilo y alineación personalizados.

    Args:
    doc (Document): El documento al que se añadirá el párrafo.
    text (str): El texto del párrafo.
    style (str, optional): El estilo del párrafo (por defecto usa 'Normal').
    alignment (WD_ALIGN_PARAGRAPH, optional): Alineación del texto (por defecto JUSTIFY).
    """
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = doc.styles[style]
    paragraph.alignment = alignment


def add_table_of_contents(doc: Document, font_size: int = 8):
    """
    Agrega una tabla de contenidos al documento con formato predefinido.

    Nota: La tabla de contenidos requiere que el documento sea abierto en Word para actualizarla manualmente.

    Args:
    doc (Document): El documento al que se añadirá la tabla de contenidos.
    font_size (int, optional): Tamaño de la fuente para todo el documento (por defecto 8 puntos).
    """
    # Configurar estilo general
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(font_size)

    # Título de la tabla de contenidos
    para = doc.add_paragraph("Tabla de Contenidos")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.name = 'Century Gothic'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 32, 96)

    # Párrafo para insertar la tabla de contenidos
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.font.name = 'Century Gothic'
    run.font.size = Pt(font_size)

    # Configurar campo TOC
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # Agregar elementos al párrafo
    run_element = run._r
    run_element.append(fldChar_begin)
    run_element.append(instrText)
    run_element.append(fldChar_separate)
    run_element.append(fldChar_end)

    # Estilo para el párrafo
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Nueva página
    doc.add_page_break()


def add_bullet_points(doc: Document, bullet_points: list, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY):
    """
    Agrega una lista con bullets al documento.

    Args:
    doc (Document): El documento al que se añadirá la lista.
    bullet_points (list): Lista de textos que irán en cada punto.
    alignment (WD_ALIGN_PARAGRAPH, optional): Alineación del texto (por defecto JUSTIFY).
    """
    for point in bullet_points:
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(point)
        paragraph.alignment = alignment

        # Configurar alineación interna (justificado si se elige)
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both' if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY else alignment)
        pPr.append(jc)

def add_numbered_list(doc: Document, items: list, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
    """
    Agrega una lista numerada al documento.

    Args:
    doc (Document): El documento al que se añadirá la lista numerada.
    items (list): Lista de textos que irán en cada elemento numerado.
    alignment (WD_ALIGN_PARAGRAPH, optional): Alineación del texto (por defecto LEFT).
    """
    for item in items:
        paragraph = doc.add_paragraph(style='List Number')
        run = paragraph.add_run(item)
        paragraph.alignment = alignment


def add_footer(doc: Document, text: str, font_size: int = 9, font_color: RGBColor = RGBColor(128, 128, 128)):
    """
    Agrega texto al pie de página del documento.

    Args:
    doc (Document): El documento al que se añadirá el texto del pie de página.
    text (str): El texto a añadir.
    font_size (int, optional): Tamaño de la fuente (por defecto 9 puntos).
    font_color (RGBColor, optional): Color del texto (por defecto gris).
    """
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.color.rgb = font_color
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_header(doc: Document, text: str, font_size: int = 10, font_color: RGBColor = RGBColor(0, 0, 0)):
    """
    Agrega texto al encabezado del documento.

    Args:
    doc (Document): El documento al que se añadirá el texto del encabezado.
    text (str): El texto a añadir.
    font_size (int, optional): Tamaño de la fuente (por defecto 10 puntos).
    font_color (RGBColor, optional): Color del texto (por defecto negro).
    """
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        run.font.color.rgb = font_color
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_footnote(doc: Document, text: str, font_size: int = 8):
    """
    Agrega una nota de pie de página al documento.

    Nota: Las notas de pie no son nativas en `python-docx` y requieren manipulación avanzada con XML.

    Args:
    doc (Document): El documento al que se añadirá la nota de pie de página.
    text (str): El texto de la nota de pie.
    font_size (int, optional): Tamaño de la fuente de la nota de pie (por defecto 8 puntos).
    """
    # Crear párrafo para la referencia en el texto
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("1")
    run.font.superscript = True  # Convertir a superíndice
    run.font.size = Pt(font_size)

    # Agregar contenido de la nota de pie
    endnotes = doc.add_paragraph()
    endnotes.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_note = endnotes.add_run(f"1. {text}")
    run_note.font.size = Pt(font_size)
    run_note.font.color.rgb = RGBColor(128, 128, 128)

def add_image_with_caption(doc: Document, image_path: str, caption_text: str = None, width: Inches = None, height: Inches = None, alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER):
    """
    Agrega una imagen al documento con opciones de tamaño, alineación y una leyenda o fuente opcional debajo de la imagen.

    Args:
    doc (Document): El documento al que se añadirá la imagen.
    image_path (str): Ruta de la imagen a insertar.
    caption_text (str, optional): El texto de la leyenda o fuente para la imagen.
    width (Inches, optional): Ancho de la imagen (por defecto mantiene el tamaño original).
    height (Inches, optional): Alto de la imagen (por defecto mantiene el tamaño original).
    alignment (WD_ALIGN_PARAGRAPH, optional): Alineación de la imagen (por defecto CENTER).
    """
    # Agregar la imagen al documento
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment

    run = paragraph.add_run()
    run.add_picture(image_path, width=width, height=height)

    # Si se proporciona un texto de leyenda o fuente, agregarlo
    if caption_text:
        caption_paragraph = doc.add_paragraph(caption_text, style='Custom Caption')
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Ajustar formato del párrafo de la leyenda
        caption_paragraph_format = caption_paragraph.paragraph_format
        caption_paragraph_format.keep_together = True  # Mantener en la misma página que la imagen
        caption_paragraph_format.keep_with_next = False  # No unir con el contenido siguiente
        caption_paragraph_format.left_indent = Cm(0.75)  # Sangría izquierda
        caption_paragraph_format.space_before = Pt(6)  # Espacio antes del párrafo
        caption_paragraph_format.space_after = Pt(6)  # Espacio después del párrafo


        